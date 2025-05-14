import sys
import os
import platform
from datetime import datetime, timezone, timedelta
import re
import json
from urllib.parse import quote
import math # Added for entropy calculation
from collections import Counter # Added for entropy calculation

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QFileDialog, QLabel, QTextEdit, QTreeWidget,
    QTreeWidgetItem, QHeaderView, QStatusBar, QProgressBar,
    QLineEdit, QStyledItemDelegate, QTreeWidgetItemIterator,
    QDialog, QDialogButtonBox, QMenu, QStyle, QStyleFactory, QCheckBox,
    QGroupBox, QListWidget, QListWidgetItem, QDockWidget,
    QStackedWidget, QFrame, QSizePolicy
)
from PyQt6.QtGui import QIcon

from PyQt6.QtCore import Qt, QSize, QTimer, QRect, QUrl
from PyQt6.QtGui import (
    QColor, QPainter, QAction, QPalette, QIcon, QPixmap, QMovie,
    QDesktopServices, QCursor
)


try:
    from PIL import Image, UnidentifiedImageError
    from PIL.ExifTags import TAGS, GPSTAGS
except ImportError:
    print("Warning: Pillow library not found. Image processing will fail.")
    Image = None
    UnidentifiedImageError = None
    TAGS = {}
    GPSTAGS = {}

try:
    import PyPDF2
except ImportError:
    print("Warning: PyPDF2 library not found. PDF processing will fail.")
    PyPDF2 = None

try:
    from docx import Document
except ImportError:
    print("Warning: python-docx library not found. DOCX processing will fail.")
    Document = None

try:
    import mutagen
except ImportError:
    print("Warning: mutagen library not found. Audio/Video processing will fail.")
    mutagen = None

try:
    import magic
except ImportError:
    print("Warning: python-magic library not found. File type verification will be limited.")
    magic = None

try:
    from stegano import lsb # Added for stegano detection
except ImportError:
    print("Warning: stegano library not found. Steganography detection (LSB) will be unavailable.")
    lsb = None


def parse_pdf_date(date_str):
    """Parses PDF date strings into datetime objects."""
    if not isinstance(date_str, str):
        return None

    match = re.match(r"D:(\d{4})(\d{2})(\d{2})(\d{2})(\d{2})(\d{2})([Zz]|([+\-])(\d{2})'(\d{2})')?", date_str)
    if match:
        try:
            year, month, day, hour, minute, second = map(int, match.groups()[:6])
            tz_part = match.group(7)
            offset_sign = match.group(9)
            offset_h = match.group(10)
            offset_m = match.group(11)

            if not (1 <= month <= 12 and 1 <= day <= 31 and 0 <= hour <= 23 and 0 <= minute <= 59 and 0 <= second <= 59):
                print(f"Warning: Invalid date/time components in PDF date '{date_str}'")
                return None

            dt = datetime(year, month, day, hour, minute, second)

            if tz_part and tz_part.lower() == 'z':
                dt = dt.replace(tzinfo=timezone.utc)
            elif offset_sign and offset_h and offset_m:
                offset_h = int(offset_h)
                offset_m = int(offset_m)
                delta = timedelta(hours=offset_h, minutes=offset_m)
                if offset_sign == '-':
                    delta = -delta
                tz = timezone(delta)
                dt = dt.replace(tzinfo=tz)

            return dt
        except (ValueError, TypeError) as e:
            print(f"Warning: Error parsing PDF date string '{date_str}': {e}")
            return None
    return None
def emoji_to_icon(emoji, size=32):
    label = QLabel()
    label.setText(emoji)
    label.setStyleSheet(f"font-size: {size}px;")
    label.resize(size, size)


    pixmap = QPixmap(label.size())
    pixmap.fill(Qt.GlobalColor.transparent)
    label.render(pixmap)

    return QIcon(pixmap)

def get_icon(icon_name):
    """Get system icon or fallback to theme icon."""
    if platform.system() == 'Windows':
        icons = {
            'file': '📄',
            'image': '🖼️',
            'pdf': '📕',
            'docx': '📝',
            'audio': '🎵',
            'video': '🎬',
            'folder': '📁',
            'search': '🔍',
            'analyze': '⚙️',
            'clear': '❌',
            'anomaly': '⚠️',
            'issue': '❓',
            'theme': '🌙',
            'export': '💾',
            'map': '🌍',
            'preview': '👁️'
        }
        emoji = icons.get(icon_name, '')
        return emoji_to_icon(emoji) if emoji else QIcon()

    else:

        theme_icons = {
            'file': QStyle.StandardPixmap.SP_FileIcon,
            'folder': QStyle.StandardPixmap.SP_DirIcon,
            'search': QStyle.StandardPixmap.SP_FileDialogContentsView,
            'clear': QStyle.StandardPixmap.SP_DialogCloseButton,
            'analyze': QStyle.StandardPixmap.SP_MediaPlay,
            'export': QStyle.StandardPixmap.SP_DialogSaveButton,
            'map': QStyle.StandardPixmap.SP_ArrowForward,
            'preview': QStyle.StandardPixmap.SP_FileDialogDetailedView
        }
        if icon_name in theme_icons:
            return QApplication.style().standardIcon(theme_icons[icon_name])

        return QIcon()


class HighlightDelegate(QStyledItemDelegate):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.highlight_color = QColor(255, 255, 0, 120)
        self.search_text = ""
        self.case_sensitive = False

    def set_search_text(self, text):
        self.search_text = text.strip()
        if self.parent() and hasattr(self.parent(), 'viewport'):
            self.parent().viewport().update()

    def set_case_sensitive(self, sensitive):
        if self.case_sensitive != sensitive:
            self.case_sensitive = sensitive
            if self.parent() and hasattr(self.parent(), 'viewport'):
                self.parent().viewport().update()

    def paint(self, painter, option, index):
        super().paint(painter, option, index)

        if not self.search_text or index.column() < 0:
            return

        text = index.data(Qt.ItemDataRole.DisplayRole)
        if not text or not isinstance(text, str):
            return

        is_selected = option.state & QStyle.StateFlag.State_Selected
        compare_text = text if self.case_sensitive else text.lower()
        compare_search = self.search_text if self.case_sensitive else self.search_text.lower()

        if compare_search in compare_text and not is_selected:
            fm = option.fontMetrics
            text_rect = option.rect.adjusted(2, 1, -2, -1)

            start_pos = compare_text.find(compare_search)
            while start_pos != -1:
                before_text = text[:start_pos]
                matched_text = text[start_pos : start_pos + len(compare_search)]

                x_start = fm.horizontalAdvance(before_text)
                match_width = fm.horizontalAdvance(matched_text)

                highlight_rect = QRect(text_rect.x() + x_start, text_rect.y(), match_width, text_rect.height())

                painter.save()
                painter.setPen(Qt.PenStyle.NoPen)
                painter.setBrush(self.highlight_color)
                painter.drawRect(highlight_rect)
                painter.restore()

                start_pos = compare_text.find(compare_search, start_pos + 1)


class MetadataTreeWidget(QTreeWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.customContextMenuRequested.connect(self.open_context_menu)
        self.file_path_map = {}


    def open_context_menu(self, pos):
        item = self.itemAt(pos)
        if item is None:
            return

        index = self.indexAt(pos)
        menu = QMenu(self)
        is_top_level = item.parent() is None
        is_property_col = index.column() == 0
        is_value_col = index.column() == 1
        prop_text = item.text(0)
        value_text = item.text(1) if item.columnCount() > 1 else ""


        if not is_top_level and prop_text == "GPS Info" and is_property_col:
            lat = None
            lon = None
            coordinate_value_text = None


            for i in range(item.childCount()):
                child = item.child(i)
                if child.text(0) == "GPSPosition":
                    coordinate_value_text = child.text(1)
                    break

            if coordinate_value_text:

                potential_gps_text = re.sub(r'\s*\(Altitude:.*\)', '', coordinate_value_text).strip()
                match = re.match(r'^([-+]?\d{1,3}(?:\.\d+)?)\s*,\s*([-+]?\d{1,3}(?:\.\d+)?)$', potential_gps_text)
                if match:
                    try:
                        lat_val = float(match.group(1))
                        lon_val = float(match.group(2))
                        if -90 <= lat_val <= 90 and -180 <= lon_val <= 180:
                            lat = lat_val
                            lon = lon_val
                    except ValueError:
                        pass

            if lat is not None and lon is not None:
                map_action = QAction(get_icon('map'), f"View Location ({lat:.4f}, {lon:.4f}) on Map", self)

                map_url = QUrl(f"https://www.google.com/maps?q={lat},{lon}")
                map_action.triggered.connect(lambda checked=False, url=map_url: QDesktopServices.openUrl(url))
                menu.addAction(map_action)
                menu.addSeparator()


        if is_top_level and is_property_col:
            preview_action = QAction(get_icon('preview'), "Show Preview", self)
            preview_action.triggered.connect(lambda: self.show_file_preview(item))
            menu.addAction(preview_action)
            menu.addSeparator()

            menu.addAction(f"Copy Filename '{prop_text}'", lambda: self.copy_cell_text(item, 0))
            menu.addAction("Copy All Metadata for this File", lambda: self.copy_all_metadata(item))
            menu.addSeparator()

            export_action = QAction(get_icon('export'), "Export Metadata as JSON", self)
            export_action.triggered.connect(lambda: self.export_metadata(item))
            menu.addAction(export_action)

        elif not is_top_level:

            if prop_text != "GPS Info":
                if is_property_col:
                     menu.addAction(f"Copy Property '{prop_text}'", lambda: self.copy_cell_text(item, 0))
                     if value_text:
                         menu.addAction(f"Copy Value '{value_text}'", lambda: self.copy_cell_text(item, 1))
                         menu.addAction(f"Copy '{prop_text}: {value_text}'", lambda: self.copy_key_value(item))

                elif is_value_col and value_text:
                     menu.addAction(f"Copy Value '{value_text}'", lambda: self.copy_cell_text(item, 1))
                     menu.addAction(f"Copy '{prop_text}: {value_text}'", lambda: self.copy_key_value(item))

            elif prop_text == "GPS Info" and item.parent():
                 parent_file_item = item.parent()
                 parent_filename = parent_file_item.text(0)
                 menu.addAction(f"Copy All Metadata for '{parent_filename}'", lambda: self.copy_all_metadata(parent_file_item))


        if menu.actions():
            menu.exec(self.viewport().mapToGlobal(pos))


    def show_file_preview(self, item):
        """Shows preview for the selected file"""

        parent = self.parent()
        while parent is not None and not isinstance(parent, MetadataAnalyzerApp):
            parent = parent.parent()

        if parent and hasattr(parent, 'update_preview'):

            filename = item.text(0)
            file_path = self.file_path_map.get(filename)
            if file_path:
                parent.update_preview(file_path)

    def copy_all_metadata(self, item):
        """Copies the filename and all child metadata recursively to the clipboard."""
        texts = [f"--- Metadata for {item.text(0)} ---"]
        self._collect_child_metadata(item, texts, indent="")
        QApplication.clipboard().setText("\n".join(texts))

    def _collect_child_metadata(self, item, texts_list, indent):
        """Recursive helper to collect metadata text."""
        for i in range(item.childCount()):
            child = item.child(i)
            prop = child.text(0)
            val = child.text(1) if child.text(1) else ""
            if val or child.childCount() == 0:
                texts_list.append(f"{indent}{prop}: {val}")
            else:
                texts_list.append(f"{indent}{prop}:")
            if child.childCount() > 0:
                self._collect_child_metadata(child, texts_list, indent + "  ")

    def copy_cell_text(self, item, column):
        """Copies the text of a specific cell to the clipboard."""
        QApplication.clipboard().setText(item.text(column))

    def copy_key_value(self, item):
        """Copies the 'Key: Value' pair to the clipboard."""
        QApplication.clipboard().setText(f"{item.text(0)}: {item.text(1)}")

    def export_metadata(self, item):
        """Exports metadata as JSON to a file."""
        metadata = {}
        self._collect_metadata_dict(item, metadata)

        file_path, _ = QFileDialog.getSaveFileName(
            self, "Export Metadata",
            f"{item.text(0)}_metadata.json",
            "JSON Files (*.json);;All Files (*)"
        )

        if file_path:
            try:
                with open(file_path, 'w') as f:
                    json.dump(metadata, f, indent=4)

                parent_app = self.window()
                if hasattr(parent_app, 'status_bar'):
                    parent_app.status_bar.showMessage(f"Metadata exported to {file_path}")
            except Exception as e:
                parent_app = self.window()
                if hasattr(parent_app, 'status_bar'):
                    parent_app.status_bar.showMessage(f"Error exporting metadata: {e}")


    def _collect_metadata_dict(self, item, metadata_dict):
        """Recursive helper to collect metadata as dictionary."""
        for i in range(item.childCount()):
            child = item.child(i)
            prop = child.text(0)
            val = child.text(1) if child.text(1) else None

            if child.childCount() > 0:
                child_dict = {}
                self._collect_metadata_dict(child, child_dict)
                metadata_dict[prop] = child_dict
            else:
                if val is not None:
                    metadata_dict[prop] = val

class PreviewWidget(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.layout = QVBoxLayout(self)
        self.setLayout(self.layout)

        self.preview_label = QLabel("No preview available")
        self.preview_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.preview_label.setStyleSheet("font-style: italic; color: #666;")
        self.layout.addWidget(self.preview_label)

        self.movie = None

    def clear_preview(self):
        """Clears the current preview."""
        if self.movie:
            self.movie.stop()
            self.movie = None
        self.preview_label.clear()
        self.preview_label.setText("No preview available")

    def set_image_preview(self, file_path):
        """Sets an image preview."""
        try:
            pixmap = QPixmap(file_path)
            if not pixmap.isNull():

                scaled_pixmap = pixmap.scaled(
                    400, 400,
                    Qt.AspectRatioMode.KeepAspectRatio,
                    Qt.TransformationMode.SmoothTransformation
                )
                self.preview_label.setPixmap(scaled_pixmap)
                self.preview_label.setText("")
            else:
                self.preview_label.setText("Image preview not available")
        except Exception as e:
            print(f"Error loading image preview: {e}")
            self.preview_label.setText("Error loading preview")

    def set_pdf_preview(self, file_path):
        """Sets a PDF preview (first page as image)."""
        self.preview_label.setText("PDF preview requires additional libraries")

    def set_video_preview(self, file_path):
        """Sets a video preview (first frame)."""
        self.preview_label.setText("Video preview requires additional libraries")

    def set_loading_animation(self):
        """Sets a loading animation."""
        self.clear_preview()
        try:

            self.movie = QMovie(":/icons/loading.gif")
            if self.movie.isValid():
                self.preview_label.setMovie(self.movie)
                self.movie.start()
            else:
                self.preview_label.setText("Loading...")

        except Exception as e:
            self.preview_label.setText("Loading...")


    def stop_loading_animation(self):
        """Stops any loading animation."""
        if self.movie:
            self.movie.stop()
            self.movie = None



class AnomaliesDialog(QDialog):
    def __init__(self, anomalies_list, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Detected Anomalies")
        self.setMinimumSize(600, 400)
        self.layout = QVBoxLayout(self)

        header = QLabel("⚠️ Potential Issues Detected")
        header.setStyleSheet("font-size: 14px; font-weight: bold; color: #d35400;")
        self.layout.addWidget(header)

        self.text_edit = QTextEdit()
        self.text_edit.setReadOnly(True)
        self.text_edit.setText("\n".join(anomalies_list) if anomalies_list else "No anomalies detected.")
        self.layout.addWidget(self.text_edit)

        self.button_box = QDialogButtonBox()
        self.ok_button = self.button_box.addButton(QDialogButtonBox.StandardButton.Ok)
        self.copy_button = QPushButton("Copy All to Clipboard")
        self.copy_button.setIcon(get_icon('export'))
        self.copy_button.clicked.connect(self.copy_all)
        self.button_box.addButton(self.copy_button, QDialogButtonBox.ButtonRole.ActionRole)
        self.layout.addWidget(self.button_box)
        self.button_box.accepted.connect(self.accept)

    def copy_all(self):
        QApplication.clipboard().setText(self.text_edit.toPlainText())
        self.copy_button.setText("Copied!")
        QTimer.singleShot(1500, lambda: self.copy_button.setText("Copy All to Clipboard"))


class LogicalIssuesDialog(QDialog):
    def __init__(self, issues_list, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Detected Logical Issues")
        self.setMinimumSize(600, 400)
        self.layout = QVBoxLayout(self)

        header = QLabel("❓ Logical Inconsistencies Detected")
        header.setStyleSheet("font-size: 14px; font-weight: bold; color: #8e44ad;")
        self.layout.addWidget(header)

        self.text_edit = QTextEdit()
        self.text_edit.setReadOnly(True)
        self.text_edit.setText("\n\n".join(issues_list) if issues_list else "No logical issues detected.")
        self.layout.addWidget(self.text_edit)

        self.button_box = QDialogButtonBox()
        self.ok_button = self.button_box.addButton(QDialogButtonBox.StandardButton.Ok)
        self.copy_button = QPushButton("Copy All to Clipboard")
        self.copy_button.setIcon(get_icon('export'))
        self.copy_button.clicked.connect(self.copy_all)
        self.button_box.addButton(self.copy_button, QDialogButtonBox.ButtonRole.ActionRole)
        self.layout.addWidget(self.button_box)
        self.button_box.accepted.connect(self.accept)

    def copy_all(self):
        QApplication.clipboard().setText(self.text_edit.toPlainText())
        self.copy_button.setText("Copied!")
        QTimer.singleShot(1500, lambda: self.copy_button.setText("Copy All to Clipboard"))


class MetadataAnalyzerApp(QMainWindow):

    COLOR_PALETTE = {
        "darkest": "#0D1321",
        "darker": "#1D2D44",
        "dark": "#3E5C76",
        "light": "#748CAB",
        "lightest": "#F0EBD8"
    }
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Metadata Analyzer Pro")
        self.setMinimumSize(QSize(1200, 800))
        self.file_paths = []
        self.anomalies = []
        self.logical_issues = []
        self.current_file_index = 0
        self.current_file_path = None
        self.processing_timer = QTimer(self)
        self.processing_timer.setInterval(10)
        self.processing_timer.timeout.connect(self.process_next_file)


        self.init_ui()


        if hasattr(self, 'tree'):

            self.tree.file_path_map = {}


        self.apply_light_theme()


    def init_ui(self):

        self.main_widget = QWidget()
        self.setCentralWidget(self.main_widget)
        self.main_layout = QHBoxLayout(self.main_widget)
        self.main_layout.setContentsMargins(0, 0, 0, 0)
        self.main_layout.setSpacing(0)


        self.create_sidebar()


        self.create_main_content()


        self.create_status_bar()
        self.status_bar.showMessage("Ready")

    def create_sidebar(self):
        """Creates the sidebar navigation."""
        self.sidebar = QDockWidget("Navigation", self)
        self.sidebar.setFeatures(QDockWidget.DockWidgetFeature.NoDockWidgetFeatures)
        self.sidebar.setTitleBarWidget(QWidget())
        self.addDockWidget(Qt.DockWidgetArea.LeftDockWidgetArea, self.sidebar)

        sidebar_widget = QWidget()
        sidebar_layout = QVBoxLayout(sidebar_widget)
        sidebar_layout.setContentsMargins(5, 10, 5, 10)
        sidebar_layout.setSpacing(10)


        file_group = QGroupBox("Files")
        file_layout = QVBoxLayout(file_group)

        self.select_files_btn = QPushButton(" Select Files")
        self.select_files_btn.setIcon(get_icon('file'))
        self.select_files_btn.clicked.connect(self.select_files)

        self.select_folder_btn = QPushButton(" Select Folder")
        self.select_folder_btn.setIcon(get_icon('folder'))
        self.select_folder_btn.clicked.connect(self.select_folder)

        file_layout.addWidget(self.select_files_btn)
        file_layout.addWidget(self.select_folder_btn)
        file_layout.addStretch()


        analysis_group = QGroupBox("Analysis")
        analysis_layout = QVBoxLayout(analysis_group)

        self.analyze_btn = QPushButton(" Analyze Files")
        self.analyze_btn.setIcon(get_icon('analyze'))
        self.analyze_btn.clicked.connect(self.start_analysis)

        self.clear_btn = QPushButton(" Clear Results")
        self.clear_btn.setIcon(get_icon('clear'))
        self.clear_btn.clicked.connect(self.clear_results)

        # Added checkbox for steganography detection
        self.check_stegano_checkbox = QCheckBox("Check for Steganography")
        analysis_layout.addWidget(self.analyze_btn)
        analysis_layout.addWidget(self.clear_btn)
        analysis_layout.addWidget(self.check_stegano_checkbox) # Added the checkbox here
        analysis_layout.addStretch()


        view_group = QGroupBox("View")
        view_layout = QVBoxLayout(view_group)

        self.anomalies_btn = QPushButton(" Show Anomalies")
        self.anomalies_btn.setIcon(get_icon('anomaly'))
        self.anomalies_btn.clicked.connect(self.show_anomalies)
        self.anomalies_btn.setEnabled(False)

        self.logical_issues_btn = QPushButton(" Show Logical Issues")
        self.logical_issues_btn.setIcon(get_icon('issue'))
        self.logical_issues_btn.clicked.connect(self.show_logical_issues)
        self.logical_issues_btn.setEnabled(False)

        self.theme_toggle_btn = QPushButton(" Switch to Dark Theme")
        self.theme_toggle_btn.setIcon(get_icon('theme'))
        self.theme_toggle_btn.clicked.connect(self.toggle_theme)

        view_layout.addWidget(self.anomalies_btn)
        view_layout.addWidget(self.logical_issues_btn)
        view_layout.addWidget(self.theme_toggle_btn)
        view_layout.addStretch()


        sidebar_layout.addWidget(file_group)
        sidebar_layout.addWidget(analysis_group)
        sidebar_layout.addWidget(view_group)
        sidebar_layout.addStretch()

        self.sidebar.setWidget(sidebar_widget)
        self.sidebar.setMaximumWidth(220)

    def create_main_content(self):
        """Creates the main content area with tabs."""
        self.content_widget = QWidget()
        self.content_layout = QVBoxLayout(self.content_widget)
        self.content_layout.setContentsMargins(10, 10, 10, 10)
        self.content_layout.setSpacing(10)


        search_bar = QHBoxLayout()
        search_bar.addWidget(QLabel("Search:"))

        self.search_box = QLineEdit()
        self.search_box.setPlaceholderText("Search metadata...")
        self.search_box.setClearButtonEnabled(True)
        self.search_box.textChanged.connect(self.filter_metadata)

        self.case_sensitive_checkbox = QCheckBox("Case Sensitive")
        self.case_sensitive_checkbox.stateChanged.connect(self.filter_metadata)

        search_bar.addWidget(self.search_box)
        search_bar.addWidget(self.case_sensitive_checkbox)
        self.content_layout.addLayout(search_bar)


        self.stacked_widget = QStackedWidget()


        self.metadata_view = QWidget()
        metadata_layout = QHBoxLayout(self.metadata_view)


        self.tree = MetadataTreeWidget()
        self.tree.setHeaderLabels(["Property", "Value"])
        self.tree.header().setSectionResizeMode(0, QHeaderView.ResizeMode.Interactive)
        self.tree.header().setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)
        self.tree.setAlternatingRowColors(True)
        self.tree.setColumnWidth(0, 300)



        self.preview_panel = QDockWidget("Preview", self)
        self.preview_panel.setFeatures(QDockWidget.DockWidgetFeature.NoDockWidgetFeatures)
        self.preview_panel.setTitleBarWidget(QWidget())
        self.addDockWidget(Qt.DockWidgetArea.RightDockWidgetArea, self.preview_panel)

        self.preview_widget = PreviewWidget()
        self.preview_panel.setWidget(self.preview_widget)
        self.preview_panel.setMaximumWidth(300)

        metadata_layout.addWidget(self.tree)
        self.stacked_widget.addWidget(self.metadata_view)

        self.content_layout.addWidget(self.stacked_widget)
        self.main_layout.addWidget(self.content_widget)


        self.highlight_delegate = HighlightDelegate(self.tree)
        self.tree.setItemDelegate(self.highlight_delegate)

    def create_status_bar(self):
        self.status_bar = QStatusBar()
        self.setStatusBar(self.status_bar)

        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        self.progress_bar.setFixedWidth(200)
        self.progress_bar.setVisible(False)

        self.status_bar.addPermanentWidget(self.progress_bar)

    def check_file_signature_mismatch(self, file_path):
        """Detects mismatches between file extension and actual content type."""
        if not magic:
            return None

        try:
            true_type = magic.from_file(file_path)
            ext = os.path.splitext(file_path)[1].lower()

            red_flags = {
                (".jpg", "PDF"): "⚠️ FAKE IMAGE: This is actually a PDF document!",
                (".jpg", "executable"): "🚨 MALWARE: This 'image' is an executable!",
                (".pdf", "executable"): "🚨 MALWARE: This PDF is an executable!",
                (".docx", "executable"): "🚨 MALWARE: This document is an executable!",
                (".exe", "text"): "⚠️ SUSPICIOUS: Executable masquerading as text",
                (".zip", "executable"): "🚨 MALWARE: Archive is actually an executable!",
                (".png", "executable"): "🚨 MALWARE: This image is an executable!"
            }

            for (expected_ext, type_keyword), warning in red_flags.items():
                if ext == expected_ext and type_keyword in true_type:
                    return warning


            if "executable" in true_type and not ext.lower() in ('.exe', '.dll', '.bat', '.ps1', '.sh'):
                return f"⚠️ SUSPICIOUS: {ext} file is actually an executable!"

        except Exception as e:
            print(f"Error checking file signature: {e}")

        return None

    def check_suspicious_authors(self, metadata, file_type):
        """Flags suspicious author information based on file type."""
        warnings = []


        if file_type == "docx":
            author = metadata.get("Last Modified By", "")
            if "admin" in author.lower():
                warnings.append(f"SUSPECT USER: Last modified by '{author}'")


        elif file_type == "pdf":
            creator = metadata.get("Creator", "")
            producer = metadata.get("Producer", "")
            if "photoshop" in creator.lower():
                warnings.append(f"UNEXPECTED CREATOR: Created with Photoshop ({creator})")
            if "crack" in producer.lower() or "keygen" in producer.lower():
                warnings.append(f"SUSPICIOUS PRODUCER: {producer}")


        elif file_type == "image":
            serial = metadata.get("SerialNumber", "")
            if serial and len(serial) > 10:
                warnings.append(f"CAMERA ID: {serial}")

        return warnings

    def check_file_size_anomalies(self, file_path, file_stats):
        """Detects suspicious file size patterns."""
        warnings = []
        base_name = os.path.basename(file_path)


        if file_stats.st_size == 0:
            warnings.append("🚨 EMPTY FILE: Possible malware placeholder or incomplete transfer")


        ext = os.path.splitext(file_path)[1].lower()
        if ext in ('.jpg', '.png', '.pdf', '.docx') and file_stats.st_size < 512:
            warnings.append(f"⚠️ SUSPICIOUSLY SMALL: Only {self.format_size(file_stats.st_size)} for a {ext.upper()} file")

        return warnings

    def check_steganography(self, file_path, file_item):
        base_name = os.path.basename(file_path)
        stegano_root = QTreeWidgetItem(file_item, ["Steganography Analysis"])
        has_stegano_findings = False

        if lsb and file_path.lower().endswith(('.png', '.bmp')):
            try:
                hidden_message = None
                try:
                     hidden_message = lsb.reveal(file_path)
                except Exception:
                     pass

                if hidden_message and hidden_message.strip():
                    warning = f"🔍 STEGANOGRAPHY (LSB): Possible hidden data detected."
                    warn_item = QTreeWidgetItem(stegano_root, ["⚠️ LSB Detection", warning])
                    warn_item.setForeground(0, QColor("orange"))
                    self.anomalies.append(f"File '{base_name}': {warning}")
                    has_stegano_findings = True
                else:
                    QTreeWidgetItem(stegano_root, ["LSB Detection", "No easily extractable LSB data found."])

            except PermissionError as e:
                perm_error_msg = f"[Permission Denied] Could not check LSB: {e}"
                QTreeWidgetItem(stegano_root, ["LSB Detection Error", perm_error_msg]).setForeground(1, QColor("gray"))
                print(f"Warning: Permission error during LSB check for {base_name}: {e}")

            except Exception as e:
                QTreeWidgetItem(stegano_root, ["LSB Detection Error", f"Failed to check Lsb: {e}"]).setForeground(1, QColor("red"))
                self.anomalies.append(f"File '{base_name}': Error checking LSB steganography - {e}")
                has_stegano_findings = True


        elif lsb and not file_path.lower().endswith(('.png', '.bmp')):
             QTreeWidgetItem(stegano_root, ["LSB Detection", "Skipped (only supports PNG and BMP)."])

        elif not lsb:
            QTreeWidgetItem(stegano_root, ["LSB Detection", "Stegano library not found."]).setForeground(1, QColor("orange"))


        if file_path.lower().endswith(('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff')):
            try:
                entropy = self.calculate_entropy(file_path)
                entropy_threshold = 7.8

                # Create the Entropy item without forcing a specific color
                entropy_item = QTreeWidgetItem(stegano_root, ["Entropy", f"{entropy:.4f}"])
                # Removed: entropy_item.setForeground(0, entropy_color)


                if entropy > entropy_threshold:
                    warning = f"⚠️ ENTROPY ANOMALY: High entropy ({entropy:.4f}) detected - suggests possible hidden data or encryption."
                    warn_item = QTreeWidgetItem(stegano_root, ["⚠️ High Entropy", warning])
                    warn_item.setForeground(0, QColor("orange"))
                    self.anomalies.append(f"File '{base_name}': {warning}")
                    has_stegano_findings = True
                else:
                     QTreeWidgetItem(stegano_root, ["Entropy Status", "Entropy within expected range."])

            except FileNotFoundError:
                 QTreeWidgetItem(stegano_root, ["Entropy Calculation Error", "File not found for entropy calculation."]).setForeground(1, QColor("red"))
                 self.anomalies.append(f"File '{base_name}': File not found during entropy calculation.")
                 has_stegano_findings = True
            except Exception as e:
                QTreeWidgetItem(stegano_root, ["Entropy Calculation Error", f"Failed to calculate entropy: {e}"]).setForeground(1, QColor("red"))
                self.anomalies.append(f"File '{base_name}': Error calculating entropy - {e}")
                has_stegano_findings = True
        else:
            QTreeWidgetItem(stegano_root, ["Entropy Analysis", "Skipped (only applicable to image files)."])


        if stegano_root.childCount() == 0:
             QTreeWidgetItem(stegano_root, ["Status", "Analysis performed, no steganography indicators found."])
        elif not has_stegano_findings and stegano_root.childCount() > 0:
             child_texts = [stegano_root.child(i).text(0) for i in range(stegano_root.childCount())]
             if not any("⚠️" in text or "🚨" in text or "Error" in text for text in child_texts):
                 QTreeWidgetItem(stegano_root, ["Status", "Analysis performed, no steganography anomalies detected."])
    def calculate_entropy(self, file_path):
        try:
            with open(file_path, 'rb') as f:
                byte_data = f.read()
        except Exception as e:
            print(f"Error reading file for entropy calculation: {e}")
            return 0.0

        if not byte_data:
            return 0.0

        byte_counts = Counter(byte_data)
        total_bytes = len(byte_data)

        entropy = 0.0
        for count in byte_counts.values():
            if count > 0:
                probability = count / total_bytes
                entropy -= probability * math.log2(probability)

        return entropy
    def toggle_theme(self):
        if self.theme_toggle_btn.text().startswith(" Switch to Dark"):
            self.theme_toggle_btn.setText(" Switch to Light Theme")
            self.theme_toggle_btn.setIcon(get_icon('theme'))
            self.apply_dark_theme()
        else:
            self.theme_toggle_btn.setText(" Switch to Dark Theme")
            self.theme_toggle_btn.setIcon(get_icon('theme'))
            self.apply_light_theme()
    def apply_light_theme(self):
        self.setStyleSheet(f"""
        QMainWindow, QDialog {{
            background-color: {self.COLOR_PALETTE['lightest']};
        }}
        QWidget {{
            font-family: 'Segoe UI', sans-serif;
            font-size: 11px;
            color: {self.COLOR_PALETTE['darkest']};
        }}
        QPushButton {{
            padding: 8px 15px;
            border: 1px solid {self.COLOR_PALETTE['light']};
            border-radius: 4px;
            background-color: {self.COLOR_PALETTE['lightest']};
            color: {self.COLOR_PALETTE['darkest']};
            text-align: left;
        }}
        QPushButton:hover {{
            background-color: {self.COLOR_PALETTE['light']};
            border-color: {self.COLOR_PALETTE['dark']};
        }}
        QPushButton:pressed {{
            background-color: {self.COLOR_PALETTE['dark']};
            color: {self.COLOR_PALETTE['lightest']};
        }}
        QPushButton:disabled {{
            background-color: #f0f0f0;
            color: #a0a0a0;
            border-color: #d0d0d0;
        }}
        QLineEdit {{
            padding: 7px;
            border: 1px solid {self.COLOR_PALETTE['light']};
            border-radius: 4px;
            background-color: white;
            color: black;
        }}
        QTreeWidget {{
            background-color: white;
            border: 1px solid {self.COLOR_PALETTE['light']};
            border-radius: 4px;
            padding: 5px;
            color: black;
            alternate-background-color: #f9f9f9;
        }}
        QTreeWidget::item {{
            padding: 3px;
        }}
        QTreeWidget::item:selected {{
            background-color: {self.COLOR_PALETTE['dark']};
            color: white;
        }}
        QHeaderView::section {{
            background-color: {self.COLOR_PALETTE['light']};
            padding: 6px;
            border: none;
            border-bottom: 1px solid {self.COLOR_PALETTE['dark']};
            color: {self.COLOR_PALETTE['darkest']};
            font-weight: bold;
        }}
        QTextEdit {{
            background-color: white;
            border: 1px solid {self.COLOR_PALETTE['light']};
            border-radius: 4px;
            padding: 10px;
            color: black;
        }}
        QStatusBar {{
            border-top: 1px solid {self.COLOR_PALETTE['light']};
            background-color: {self.COLOR_PALETTE['lightest']};
        }}
        QStatusBar QLabel {{
            padding-left: 5px;
        }}
        QProgressBar {{
            border: 1px solid {self.COLOR_PALETTE['light']};
            border-radius: 4px;
            text-align: center;
            background-color: white;
            color: black;
        }}
        QProgressBar::chunk {{
            background-color: {self.COLOR_PALETTE['dark']};
            border-radius: 3px;
            margin: 1px;
        }}
        QMenu {{
            background-color: white;
            border: 1px solid {self.COLOR_PALETTE['light']};
            padding: 4px;
            color: black;
        }}
        QMenu::item {{
            padding: 5px 20px 5px 20px;
            background-color: transparent;
            color: black;
        }}
        QMenu::item:selected {{
            background-color: {self.COLOR_PALETTE['dark']};
            color: white;
        }}
        QMenu::separator {{
            height: 1px;
            background-color: {self.COLOR_PALETTE['light']};
            margin: 4px 0px;
        }}
        QCheckBox {{
            margin-left: 5px;
        }}
        QGroupBox {{
            border: 1px solid {self.COLOR_PALETTE['light']};
            border-radius: 5px;
            margin-top: 10px;
            padding-top: 15px;
            background-color: {self.COLOR_PALETTE['lightest']};
        }}
        QGroupBox::title {{
            subcontrol-origin: margin;
            left: 10px;
            padding: 0 3px;
            color: {self.COLOR_PALETTE['darkest']};
        }}
        QDockWidget {{
            background: {self.COLOR_PALETTE['lightest']};
            border: 1px solid {self.COLOR_PALETTE['light']};
        }}
        """           )

    def apply_dark_theme(self):
        self.setStyleSheet(f"""
        QMainWindow, QDialog {{
            background-color: {self.COLOR_PALETTE['darkest']};
        }}
        QWidget {{
            font-family: 'Segoe UI', sans-serif;
            font-size: 11px;
            color: {self.COLOR_PALETTE['lightest']};
        }}
        QPushButton {{
            padding: 8px 15px;
            border: 1px solid {self.COLOR_PALETTE['dark']};
            border-radius: 4px;
            background-color: {self.COLOR_PALETTE['darker']};
            color: {self.COLOR_PALETTE['lightest']};
            text-align: left;
        }}
        QPushButton:hover {{
            background-color: {self.COLOR_PALETTE['dark']};
            border-color: {self.COLOR_PALETTE['light']};
        }}
        QPushButton:pressed {{
            background-color: {self.COLOR_PALETTE['light']};
            color: {self.COLOR_PALETTE['darkest']};
        }}
        QPushButton:disabled {{
            background-color: #333333;
            color: #777777;
            border-color: #444444;
        }}
        QLineEdit {{
            padding: 7px;
            border: 1px solid {self.COLOR_PALETTE['dark']};
            border-radius: 4px;
            background-color: {self.COLOR_PALETTE['darker']};
            color: {self.COLOR_PALETTE['lightest']};
        }}
        QTreeWidget {{
            background-color: {self.COLOR_PALETTE['darker']};
            border: 1px solid {self.COLOR_PALETTE['dark']};
            border-radius: 4px;
            padding: 5px;
            color: {self.COLOR_PALETTE['lightest']};
            alternate-background-color: {self.COLOR_PALETTE['darkest']};
        }}
        QTreeWidget::item {{
            padding: 3px;
        }}
        QTreeWidget::item:selected {{
            background-color: {self.COLOR_PALETTE['light']};
            color: {self.COLOR_PALETTE['darkest']};
        }}
        QHeaderView::section {{
            background-color: {self.COLOR_PALETTE['darkest']};
            padding: 6px;
            border: none;
            border-bottom: 1px solid {self.COLOR_PALETTE['light']};
            color: {self.COLOR_PALETTE['lightest']};
            font-weight: bold;
        }}
        QTextEdit {{
            background-color: {self.COLOR_PALETTE['darker']};
            border: 1px solid {self.COLOR_PALETTE['dark']};
            border-radius: 4px;
            padding: 10px;
            color: {self.COLOR_PALETTE['lightest']};
        }}
        QStatusBar {{
            border-top: 1px solid {self.COLOR_PALETTE['dark']};
            background-color: {self.COLOR_PALETTE['darkest']};
        }}
        QStatusBar QLabel {{
            padding-left: 5px;
            color: {self.COLOR_PALETTE['lightest']};
        }}
        QProgressBar {{
            border: 1px solid {self.COLOR_PALETTE['dark']};
            border-radius: 4px;
            text-align: center;
            background-color: {self.COLOR_PALETTE['darker']};
            color: {self.COLOR_PALETTE['lightest']};
        }}
        QProgressBar::chunk {{
            background-color: {self.COLOR_PALETTE['light']};
            border-radius: 3px;
            margin: 1px;
        }}
        QMenu {{
            background-color: {self.COLOR_PALETTE['darkest']};
            border: 1px solid {self.COLOR_PALETTE['dark']};
            padding: 4px;
            color: {self.COLOR_PALETTE['lightest']};
        }}
        QMenu::item {{
            padding: 5px 20px 5px 20px;
            background-color: transparent;
            color: {self.COLOR_PALETTE['lightest']};
        }}
        QMenu::item:selected {{
            background-color: {self.COLOR_PALETTE['light']};
            color: {self.COLOR_PALETTE['darkest']};
        }}
        QMenu::separator {{
            height: 1px;
            background-color: {self.COLOR_PALETTE['dark']};
            margin: 4px 0px;
        }}
        QCheckBox {{
            margin-left: 5px;
            color: {self.COLOR_PALETTE['lightest']};
        }}
        QGroupBox {{
            border: 1px solid {self.COLOR_PALETTE['dark']};
            border-radius: 5px;
            margin-top: 10px;
            padding-top: 15px;
            background-color: {self.COLOR_PALETTE['darkest']};
        }}
        QGroupBox::title {{
            subcontrol-origin: margin;
            left: 10px;
            padding: 0 3px;
            color: {self.COLOR_PALETTE['lightest']};
        }}
        QDockWidget {{
            background: {self.COLOR_PALETTE['darkest']};
            border: 1px solid {self.COLOR_PALETTE['dark']};
        }}
        """           )

    def set_controls_enabled(self, enabled):
        self.select_files_btn.setEnabled(enabled)
        self.select_folder_btn.setEnabled(enabled)
        self.search_box.setEnabled(enabled)
        self.case_sensitive_checkbox.setEnabled(enabled)
        self.analyze_btn.setEnabled(enabled)
        self.clear_btn.setEnabled(enabled)
        self.check_stegano_checkbox.setEnabled(enabled) # Enable/disable checkbox with other controls


    def select_files(self):
        file_filter = "All Supported Files (*.jpg *.jpeg *.png *.gif *.bmp *.tiff *.pdf *.docx *.mp3 *.mp4);;"                    "Images (*.jpg *.jpeg *.png *.gif *.bmp *.tiff);;"                    "PDFs (*.pdf);;"                    "Word Documents (*.docx);;"                    "Audio Files (*.mp3);;"                    "Video Files (*.mp4);;"                    "All Files (*)"

        files, _ = QFileDialog.getOpenFileNames(self, "Select Files", "", file_filter)
        if files:
            self.file_paths = files


            if hasattr(self, 'tree'):
                self.tree.file_path_map = {os.path.basename(path): path for path in files}

            if len(files) == 1:
                self.status_bar.showMessage(f"Selected: {os.path.basename(files[0])}")
            else:
                self.status_bar.showMessage(f"Selected {len(files)} files")
            self.clear_results_ui_only()
        elif not self.file_paths:
            self.status_bar.showMessage("No files selected")
            self.clear_results_ui_only()

    def select_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "Select Folder")
        if folder:
            self.file_paths = []
            supported_extensions = (".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".pdf", ".docx", ".mp3", ".mp4")

            try:
                for entry in os.scandir(folder):
                    if entry.is_file() and entry.name.lower().endswith(supported_extensions):
                        abs_path = os.path.abspath(entry.path)
                        self.file_paths.append(abs_path)


                if hasattr(self, 'tree'):
                    self.tree.file_path_map = {os.path.basename(path): path for path in self.file_paths}

            except OSError as e:
                self.status_bar.showMessage(f"Error reading folder: {e}")
                return

            if self.file_paths:
                self.status_bar.showMessage(f"Selected folder: {folder} ({len(self.file_paths)} supported files found)")
                self.clear_results_ui_only()
            else:
                self.status_bar.showMessage(f"Selected folder: {folder} (No supported files found)")
                self.clear_results_ui_only()
        elif not self.file_paths:
            self.status_bar.showMessage("No folder selected")
            self.clear_results_ui_only()

    def start_analysis(self):
        if not self.file_paths:
            self.status_bar.showMessage("No files or folder selected to analyze.")
            return


        QApplication.setOverrideCursor(Qt.CursorShape.WaitCursor)


        self.current_file_index = 0
        self.anomalies = []
        self.logical_issues = []
        self.tree.clear()
        self.preview_widget.clear_preview()
        self.progress_bar.setValue(0)
        self.progress_bar.setVisible(True)

        self.anomalies_btn.setEnabled(False)
        self.anomalies_btn.setText(" Show Anomalies")
        self.logical_issues_btn.setEnabled(False)
        self.logical_issues_btn.setText(" Show Logical Issues")

        self.status_bar.showMessage(f"Starting analysis of {len(self.file_paths)} files...")
        self.set_controls_enabled(False)
        QApplication.processEvents()
        self.processing_timer.start()

    def process_next_file(self):
        try:
            if self.current_file_index >= len(self.file_paths):
                self.processing_timer.stop()
                self.finish_analysis()
                return

            file_path = self.file_paths[self.current_file_index]
            self.current_file_path = file_path
            base_name = os.path.basename(file_path)


            self.update_preview(file_path)

            file_item = QTreeWidgetItem(self.tree, [base_name])
            file_stats = None

            try:
                file_stats = os.stat(file_path)
                self.add_basic_file_info(file_path, file_item, file_stats)


                sig_warning = self.check_file_signature_mismatch(file_path)
                if sig_warning:
                    alert_item = QTreeWidgetItem(file_item, ["🚨 SIGNATURE MISMATCH", sig_warning])
                    alert_item.setForeground(0, QColor("red"))
                    alert_item.setForeground(1, QColor("red"))
                    self.anomalies.append(f"{base_name}: {sig_warning}")


                size_warnings = self.check_file_size_anomalies(file_path, file_stats)
                for warning in size_warnings:
                    warn_item = QTreeWidgetItem(file_item, ["⚠️ SIZE WARNING", warning])
                    warn_item.setForeground(0, QColor("orange"))
                    self.anomalies.append(f"{base_name}: {warning}")

            except Exception as e:
                self.anomalies.append(f"{base_name}: Error reading file system stats - {e}")
                fs_info_root = QTreeWidgetItem(file_item, ["File System Info"])
                QTreeWidgetItem(fs_info_root, ["Error", f"Could not read stats: {e}"]).setForeground(1, QColor("red"))
                file_stats = None

            lower_path = file_path.lower()
            try:
                file_type = None
                if Image and lower_path.endswith((".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff")):
                    file_type = "image"
                    self.process_image_exif(file_path, file_item, file_stats)

                    # Added Steganography Check for images
                    if self.check_stegano_checkbox.isChecked():
                         self.check_steganography(file_path, file_item)

                elif PyPDF2 and lower_path.endswith(".pdf"):
                    file_type = "pdf"
                    self.process_pdf(file_path, file_item, file_stats)
                elif Document and lower_path.endswith(".docx"):
                    file_type = "docx"
                    self.process_docx(file_path, file_item, file_stats)
                elif mutagen and lower_path.endswith((".mp3", ".mp4")):
                    file_type = "media"
                    self.process_media(file_path, file_item, file_stats)
                elif file_stats is not None:
                    has_specific_metadata = False
                    for i in range(file_item.childCount()):
                        child_text = file_item.child(i).text(0)
                        if child_text not in ["File System Info", "Status", "Processing Error", "Steganography Analysis"]: # Added Steganography Analysis here
                            has_specific_metadata = True
                            break
                    if not has_specific_metadata:
                        unsupported_item = QTreeWidgetItem(file_item, ["Status", "Unsupported file type or required library missing."])
                        unsupported_item.setForeground(1, QColor("orange"))


                if file_type:
                    metadata = {}
                    self.tree._collect_metadata_dict(file_item, metadata)
                    author_warnings = self.check_suspicious_authors(metadata, file_type)
                    for warning in author_warnings:
                        auth_item = QTreeWidgetItem(file_item, ["🔍 AUTHOR WARNING", warning])
                        auth_item.setForeground(0, QColor(139, 0, 139))
                        self.logical_issues.append(f"{base_name}: {warning}")

            except Exception as proc_err:
                self.anomalies.append(f"{base_name}: Unexpected error during processing - {proc_err}")
                error_item = QTreeWidgetItem(file_item, ["Processing Error", str(proc_err)])
                error_item.setForeground(1, QColor("red"))

            self.current_file_index += 1
            progress_percent = int((self.current_file_index / len(self.file_paths)) * 100)
            self.progress_bar.setValue(progress_percent)
            self.status_bar.showMessage(f"Analyzing: {base_name} ({self.current_file_index}/{len(self.file_paths)})")

        except Exception as e:
            QApplication.restoreOverrideCursor()
            self.status_bar.showMessage(f"Error during processing: {str(e)}")
            self.processing_timer.stop()

    def update_preview(self, file_path):
        """Updates the preview panel based on file type."""
        if not os.path.exists(file_path):
            self.preview_widget.clear_preview()
            return

        lower_path = file_path.lower()

        try:
            if lower_path.endswith((".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff")):
                self.preview_widget.set_image_preview(file_path)
            elif lower_path.endswith(".pdf"):
                self.preview_widget.set_pdf_preview(file_path)
            elif lower_path.endswith((".mp3", ".mp4")):
                self.preview_widget.set_video_preview(file_path)
            else:
                self.preview_widget.clear_preview()
        except Exception as e:
            print(f"Error generating preview: {e}")
            self.preview_widget.clear_preview()
            self.preview_widget.preview_label.setText("Preview unavailable")

    def preview_file(self):
        """Opens the current file in the default application."""
        if self.current_file_path and os.path.exists(self.current_file_path):
            try:
                QDesktopServices.openUrl(QUrl.fromLocalFile(self.current_file_path))
            except Exception as e:
                self.status_bar.showMessage(f"Error opening file: {e}")

    def filter_metadata(self):
        search_text = self.search_box.text()
        case_sensitive = self.case_sensitive_checkbox.isChecked()

        self.highlight_delegate.set_search_text(search_text)
        self.highlight_delegate.set_case_sensitive(case_sensitive)

        if not search_text:
            for i in range(self.tree.topLevelItemCount()):
                top_item = self.tree.topLevelItem(i)
                top_item.setHidden(False)
                top_item.setExpanded(True)
                for j in range(top_item.childCount()):
                    group_item = top_item.child(j)
                    group_name = group_item.text(0)

                    if group_name in ["File System Info", "EXIF Metadata", "PDF Metadata", "DOCX Metadata", "Media Metadata", "GPS Info", "Steganography Analysis"]: # Added Steganography Analysis here
                        group_item.setExpanded(True)
                    else:
                        group_item.setExpanded(False)
            self.tree.viewport().update()
            return

        matched_top_level_items = set()

        iterator = QTreeWidgetItemIterator(self.tree, QTreeWidgetItemIterator.IteratorFlag.All)
        while iterator.value():
            item = iterator.value()
            match_found = False
            for col in range(item.columnCount()):
                item_text = item.text(col)
                if item_text:
                    compare_item_text = item_text if case_sensitive else item_text.lower()
                    compare_search = search_text if case_sensitive else search_text.lower()
                    if compare_search in compare_item_text:
                        match_found = True
                        break

            if match_found:

                top_item_ancestor = item
                while top_item_ancestor.parent():
                    top_item_ancestor = top_item_ancestor.parent()
                matched_top_level_items.add(top_item_ancestor.text(0))


                if item.childCount() > 0:
                    item.setExpanded(True)
                parent = item.parent()
                while parent:
                    parent.setExpanded(True)
                    parent = parent.parent()

            iterator += 1


        for i in range(self.tree.topLevelItemCount()):
            top_item = self.tree.topLevelItem(i)
            if top_item.text(0) not in matched_top_level_items:
                top_item.setHidden(True)
            else:
                top_item.setHidden(False)
                top_item.setExpanded(True)

        self.tree.viewport().update()


    def finish_analysis(self):

        QApplication.restoreOverrideCursor()


        self.preview_widget.stop_loading_animation()

        total_files = len(self.file_paths)
        self.progress_bar.setValue(100)
        self.status_bar.showMessage(f"Analysis complete. Processed {total_files} files.")
        self.set_controls_enabled(True)
        QTimer.singleShot(1500, lambda: self.progress_bar.setVisible(False))

        if self.anomalies:
            self.anomalies_btn.setEnabled(True)
            self.anomalies_btn.setText(f" Show Anomalies ({len(self.anomalies)})")
        else:
            self.anomalies_btn.setEnabled(False)
            self.anomalies_btn.setText(" Show Anomalies")

        if self.logical_issues:
            self.logical_issues_btn.setEnabled(True)
            self.logical_issues_btn.setText(f" Show Logical Issues ({len(self.logical_issues)})")
        else:
            self.logical_issues_btn.setEnabled(False)
            self.logical_issues_btn.setText(" Show Logical Issues")


        for i in range(self.tree.topLevelItemCount()):
            top_item = self.tree.topLevelItem(i)
            if not top_item.isHidden():
                top_item.setExpanded(True)
                for j in range(top_item.childCount()):
                    group_item = top_item.child(j)
                    group_name = group_item.text(0)


                    if group_name in ["File System Info", "EXIF Metadata", "PDF Metadata", "DOCX Metadata", "Media Metadata", "GPS Info", "Steganography Analysis"]: # Added Steganography Analysis here
                        group_item.setExpanded(True)


        if self.search_box.text():
            self.filter_metadata()

    @staticmethod
    def format_size(size_bytes):
        if size_bytes is None:
            return "N/A"
        if size_bytes < 1024:
            return f"{size_bytes} B"
        elif size_bytes < 1024**2:
            return f"{size_bytes/1024:.1f} KiB"
        elif size_bytes < 1024**3:
            return f"{size_bytes/1024**2:.1f} MiB"
        else:
            return f"{size_bytes/1024**3:.1f} GiB"

    def add_basic_file_info(self, file_path, parent_item, stats):
        base_name = os.path.basename(file_path)
        fs_info_root = QTreeWidgetItem(parent_item, ["File System Info"])

        if stats is None:
            QTreeWidgetItem(fs_info_root, ["Error", "Could not read file system statistics."]).setForeground(1, QColor("red"))
            return

        creation_dt, modification_dt, access_dt = None, None, None
        creation_time_str, modification_time_str, access_time_str = "N/A", "N/A", "N/A"
        size_str = "N/A"
        creation_source = "ctime"

        try:
            ts_mod = stats.st_mtime
            ts_acc = stats.st_atime

            if hasattr(stats, 'st_birthtime') and stats.st_birthtime:
                ts_cre = stats.st_birthtime
                creation_source = "birthtime"
            else:
                ts_cre = stats.st_ctime
                if sys.platform != 'win32':
                    creation_source = "ctime (metadata change)"

            try:
                modification_dt = datetime.fromtimestamp(ts_mod, tz=timezone.utc).astimezone()
                access_dt = datetime.fromtimestamp(ts_acc, tz=timezone.utc).astimezone()
                creation_dt = datetime.fromtimestamp(ts_cre, tz=timezone.utc).astimezone()
                dt_format = "%Y-%m-%d %H:%M:%S %Z%z"
            except (OSError, ValueError):
                modification_dt = datetime.fromtimestamp(ts_mod)
                access_dt = datetime.fromtimestamp(ts_acc)
                creation_dt = datetime.fromtimestamp(ts_cre)
                dt_format = "%Y-%m-%d %H:%M:%S (Local?)"

            modification_time_str = modification_dt.strftime(dt_format)
            access_time_str = access_dt.strftime(dt_format)
            creation_time_str = creation_dt.strftime(dt_format) + f" (source: {creation_source})"

            file_size = stats.st_size
            size_str = self.format_size(file_size)


            if creation_dt and modification_dt and access_dt:
                issue_prefix = f"File '{base_name}' - Filesystem Time Issue:"
                tolerance = timedelta(seconds=2)

                if modification_dt < creation_dt - tolerance:
                    self.logical_issues.append(f"{issue_prefix} Modified ({modification_time_str}) significantly before Created ({creation_time_str})")

                if access_dt < creation_dt - tolerance:
                    self.logical_issues.append(f"{issue_prefix} Accessed ({access_time_str}) significantly before Created ({creation_time_str})")


            if file_size == 0:
                self.logical_issues.append(f"File '{base_name}': File size is 0 bytes (empty file).")

        except Exception as e:
            self.anomalies.append(f"{base_name}: Error processing file system stats - {e}")
            creation_time_str = modification_time_str = access_time_str = "Error processing"
            size_str = "Error processing"
            QTreeWidgetItem(fs_info_root, ["Error", f"Could not process timestamps/size: {e}"]).setForeground(1, QColor("red"))

        QTreeWidgetItem(fs_info_root, ["File Size", size_str])
        QTreeWidgetItem(fs_info_root, ["Created", creation_time_str])
        QTreeWidgetItem(fs_info_root, ["Modified", modification_time_str])
        QTreeWidgetItem(fs_info_root, ["Accessed", access_time_str])
        QTreeWidgetItem(fs_info_root, ["Extension", os.path.splitext(file_path)[1]])

    def _parse_exif_datetime(self, dt_str):
        if not dt_str or not isinstance(dt_str, str):
            return None
        try:

            return datetime.strptime(dt_str, '%Y:%m:%d %H:%M:%S')
        except ValueError:
            try:

                return datetime.strptime(dt_str.rstrip('\x00').strip(), '%Y:%m:%d %H:%M:%S')
            except ValueError:
                print(f"Warning: Could not parse EXIF datetime string: '{dt_str}'")
                return None

    def _parse_gps_info(self, exif_data):
        """Parses GPS information from EXIF data."""
        gps_info = {}
        if not exif_data or not hasattr(exif_data, 'get'):
            return gps_info


        for tag, value in exif_data.items():
            decoded = GPSTAGS.get(tag, tag)
            gps_info[decoded] = value

        latitude = None
        longitude = None
        altitude = None

        try:

            if 'GPSLatitude' in gps_info and 'GPSLatitudeRef' in gps_info:
                lat_deg, lat_min, lat_sec = gps_info['GPSLatitude']
                lat_ref = gps_info['GPSLatitudeRef']
                latitude = float(lat_deg) + float(lat_min)/60.0 + float(lat_sec)/3600.0
                if lat_ref == 'S':
                    latitude = -latitude


            if 'GPSLongitude' in gps_info and 'GPSLongitudeRef' in gps_info:
                lon_deg, lon_min, lon_sec = gps_info['GPSLongitude']
                lon_ref = gps_info['GPSLongitudeRef']
                longitude = float(lon_deg) + float(lon_min)/60.0 + float(lon_sec)/3600.0
                if lon_ref == 'W':
                    longitude = -longitude


            if 'GPSAltitude' in gps_info:
                altitude_val = gps_info['GPSAltitude']

                if isinstance(altitude_val, tuple) and len(altitude_val) == 2:
                    altitude = float(altitude_val[0]) / float(altitude_val[1]) if altitude_val[1] != 0 else 0.0
                else:
                    altitude = float(altitude_val)

                if 'GPSAltitudeRef' in gps_info and gps_info['GPSAltitudeRef'] == 1:
                    altitude = -altitude


            if latitude is not None and longitude is not None:
                gps_info['GPSLatitudeDec'] = latitude
                gps_info['GPSLongitudeDec'] = longitude
                gps_info['GPSPosition'] = f"{latitude:.6f}, {longitude:.6f}"
                if altitude is not None:
                     gps_info['GPSAltitude'] = f"{altitude:.1f}"



        except Exception as e:
            print(f"Error parsing GPS info: {e}")

            gps_info['GPSParsingError'] = str(e)

        return gps_info

    def process_image_exif(self, file_path, parent_item, stats):
        base_name = os.path.basename(file_path)
        if not Image or not TAGS or not UnidentifiedImageError:
            QTreeWidgetItem(parent_item, ["EXIF Status", "Pillow library missing or incomplete."]).setForeground(1, QColor("orange"))
            return

        encodings_to_try = ['utf-8', 'ascii', 'latin-1', 'windows-1252', 'utf-16le', 'utf-16be', 'shift_jis', 'cp437']
        exif_root = QTreeWidgetItem(parent_item, ["EXIF Metadata"])
        has_exif_content = False
        exif_values = {}

        fs_creation_dt, fs_modification_dt = None, None
        if stats:
            try:
                fs_modification_dt = datetime.fromtimestamp(stats.st_mtime)
                ts_cre = stats.st_ctime
                if hasattr(stats, 'st_birthtime') and stats.st_birthtime:
                    ts_cre = stats.st_birthtime
                fs_creation_dt = datetime.fromtimestamp(ts_cre)
            except Exception as fs_e:
                print(f"Warning: Could not extract filesystem times for comparison in {base_name}: {fs_e}")

        try:
            with Image.open(file_path) as image:

                image.load()
                exif_data = image.getexif()

                if exif_data:
                    found_tags = set()
                    for tag_id, value in exif_data.items():
                        tag_name = TAGS.get(tag_id, f"Unknown Tag ({tag_id})")
                        found_tags.add(tag_name)
                        has_exif_content = True
                        value_str = None


                        if isinstance(value, bytes):
                            successfully_decoded = False
                            for encoding in encodings_to_try:
                                try:
                                    value_str = value.decode(encoding, errors='ignore').rstrip('\x00').strip()

                                    if value_str.startswith('\ufeff'):
                                        value_str = value_str[1:]

                                    if value_str and not all(c in ('\x00', '\ufffd') for c in value_str):
                                        successfully_decoded = True
                                        break
                                except Exception:
                                    continue
                            if not successfully_decoded:
                                value_str = repr(value)
                        elif isinstance(value, str):
                            value_str = value.rstrip('\x00').strip()
                        else:
                            value_str = str(value).rstrip('\x00')

                        exif_values[tag_name] = value_str


                        max_len = 200
                        display_value = value_str if value_str is not None else ""
                        if len(display_value) > max_len:
                            display_value = display_value[:max_len] + "..."

                        QTreeWidgetItem(exif_root, [str(tag_name), display_value])


                    gps_ifd = exif_data.get_ifd(0x8825) if hasattr(exif_data, 'get_ifd') else {}
                    gps_info = self._parse_gps_info(gps_ifd)
                    if gps_info:
                        gps_root = QTreeWidgetItem(exif_root, ["GPS Info"])
                        for key, value in gps_info.items():

                            if key in ['GPSLatitudeDec', 'GPSLongitudeDec', 'GPSAltitude', 'GPSPosition', 'GPSTimestamp', 'GPSDateStamp', 'GPSProcessingMethod', 'GPSParsingError']:
                                QTreeWidgetItem(gps_root, [str(key), str(value)])

                    if 'SerialNumber' in exif_values:
                        serial_item = QTreeWidgetItem(exif_root, ["Camera Serial Number", exif_values['SerialNumber']])
                        serial_item.setForeground(0, QColor(0, 100, 0))
                        self.logical_issues.append(f"File '{base_name}': Camera serial number found - {exif_values['SerialNumber']}")


                    if 'Software' in exif_values:
                        software = exif_values['Software'].lower()
                        if 'photoshop' in software or 'editor' in software:
                            warn_item = QTreeWidgetItem(exif_root, ["⚠️ EDITING SOFTWARE", exif_values['Software']])
                            warn_item.setForeground(0, QColor("orange"))
                            self.anomalies.append(f"File '{base_name}': Edited with {exif_values['Software']}")


                    dt_orig_str = exif_values.get("DateTimeOriginal")
                    dt_digi_str = exif_values.get("DateTimeDigitized")
                    dt_orig = self._parse_exif_datetime(dt_orig_str)
                    dt_digi = self._parse_exif_datetime(dt_digi_str)

                    if dt_orig or dt_digi or fs_creation_dt or fs_modification_dt:
                        issue_prefix = f"File '{base_name}' - EXIF Time Issue:"
                        if dt_orig and dt_digi and dt_orig > dt_digi:
                            self.logical_issues.append(f"{issue_prefix} DateTimeOriginal ({dt_orig}) is after DateTimeDigitized ({dt_digi})")
                        if dt_orig and fs_modification_dt and dt_orig > fs_modification_dt + timedelta(minutes=1):
                            self.logical_issues.append(f"{issue_prefix} DateTimeOriginal ({dt_orig}) is significantly after Filesystem Modified ({fs_modification_dt.strftime('%Y-%m-%d %H:%M:%S')}) - Suggests file modification after capture.")
                        if dt_orig and fs_creation_dt and dt_orig < fs_creation_dt - timedelta(minutes=1):
                            self.logical_issues.append(f"File '{base_name}' - EXIF Note: DateTimeOriginal ({dt_orig}) is significantly before Filesystem Created ({fs_creation_dt.strftime('%Y-%m-%d %H:%M:%S')}) - May indicate copying or timestamp manipulation.")


                    required_tags = ["DateTimeOriginal", "Make", "Model"]
                    missing_tags = [tag for tag in required_tags if tag not in found_tags]
                    if missing_tags:
                        missing_item = QTreeWidgetItem(exif_root, ["⚠️ MISSING TAGS", ", ".join(missing_tags)])
                        missing_item.setForeground(0, QColor("orange"))
                        self.anomalies.append(f"File '{base_name}': Missing common EXIF tags: {', '.join(missing_tags)}")

                if not has_exif_content:
                    QTreeWidgetItem(exif_root, ["Status", "No EXIF data found."])

                    if file_path.lower().endswith(('.jpg', '.jpeg', '.png')):
                        self.anomalies.append(f"File '{base_name}': No EXIF data found in image file - may be stripped or edited")

        except FileNotFoundError:
            self.anomalies.append(f"{base_name}: File not found during EXIF processing.")
            QTreeWidgetItem(exif_root, ["Error", "File not found."]).setForeground(1, QColor("red"))
        except UnidentifiedImageError:
            self.anomalies.append(f"{base_name}: Cannot identify image file (may be corrupt or unsupported format).")
            QTreeWidgetItem(exif_root, ["Error", "Cannot identify image file."]).setForeground(1, QColor("orange"))
        except Exception as e:
            self.anomalies.append(f"{base_name}: Error processing image EXIF - {e}")
            QTreeWidgetItem(exif_root, ["Error", f"EXIF processing failed: {e}"]).setForeground(1, QColor("red"))

    def process_pdf(self, file_path, parent_item, stats):
        base_name = os.path.basename(file_path)
        if not PyPDF2:
            QTreeWidgetItem(parent_item, ["PDF Status", "PyPDF2 library missing."]).setForeground(1, QColor("orange"))
            return

        pdf_root = QTreeWidgetItem(parent_item, ["PDF Metadata"])
        has_pdf_content = False
        pdf_values = {}

        fs_creation_dt, fs_modification_dt = None, None
        if stats:
            try:
                fs_modification_dt = datetime.fromtimestamp(stats.st_mtime)
                ts_cre = stats.st_ctime
                if hasattr(stats, 'st_birthtime') and stats.st_birthtime:
                    ts_cre = stats.st_birthtime
                fs_creation_dt = datetime.fromtimestamp(ts_cre)
            except Exception as fs_e:
                print(f"Warning: Could not extract filesystem times for comparison in {base_name}: {fs_e}")

        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f, strict=False)
                meta = reader.metadata
                if meta:
                    has_pdf_content = True
                    fields = {
                        '/Title': "Title",
                        '/Author': "Author",
                        '/Subject': "Subject",
                        '/Producer': "Producer",
                        '/Creator': "Creator",
                        '/CreationDate': "Creation Date",
                        '/ModDate': "Modification Date",
                        '/Keywords': "Keywords"
                    }

                    for field_key, field_name in fields.items():
                        value = meta.get(field_key)
                        if value is not None:
                            value_str = str(value).strip()
                            pdf_values[field_name] = value

                            if field_key == '/Creator':
                                creator_item = QTreeWidgetItem(pdf_root, [field_name, value_str])
                                if "photoshop" in value_str.lower():
                                    creator_item.setForeground(1, QColor(139, 0, 139))
                                    self.logical_issues.append(f"File '{base_name}': Created with Photoshop ({value_str})")
                                elif "acrobat" not in value_str.lower():
                                    creator_item.setForeground(1, QColor("blue"))
                            else:
                                QTreeWidgetItem(pdf_root, [field_name, value_str])


                    other_meta = {k: meta[k] for k in meta if k not in fields and meta.get(k) is not None}
                    if other_meta:
                        other_root = QTreeWidgetItem(pdf_root, ["Other Metadata"])
                        for key, value in other_meta.items():
                            value_str = str(value).strip()
                            pdf_values[key] = value
                            QTreeWidgetItem(other_root, [str(key).lstrip('/'), value_str])


                    pdf_create_val = pdf_values.get("Creation Date")
                    pdf_mod_val = pdf_values.get("Modification Date")
                    pdf_create_dt = parse_pdf_date(pdf_create_val)
                    pdf_mod_dt = parse_pdf_date(pdf_mod_val)


                    fs_mod_naive = fs_modification_dt.replace(tzinfo=None) if fs_modification_dt else None
                    fs_cre_naive = fs_creation_dt.replace(tzinfo=None) if fs_creation_dt else None

                    if pdf_create_dt or pdf_mod_dt or fs_cre_naive or fs_mod_naive:
                        issue_prefix = f"File '{base_name}' - PDF Time Issue:"
                        if pdf_create_dt and pdf_mod_dt and pdf_mod_dt < pdf_create_dt:
                            time_item = QTreeWidgetItem(pdf_root, ["⚠️ TIME INCONSISTENCY", "ModDate before CreationDate"])
                            time_item.setForeground(0, QColor("orange"))
                            self.logical_issues.append(f"{issue_prefix} PDF ModDate ({pdf_mod_dt}) is before PDF CreationDate ({pdf_create_dt})")

                        pdf_create_naive = pdf_create_dt.replace(tzinfo=None) if pdf_create_dt and pdf_create_dt.tzinfo else pdf_create_dt
                        if pdf_create_naive and fs_mod_naive and pdf_create_naive > fs_mod_naive + timedelta(minutes=1):
                            time_item = QTreeWidgetItem(pdf_root, ["⚠️ TIME INCONSISTENCY", "PDF created after filesystem modified"])
                            time_item.setForeground(0, QColor("orange"))
                            self.logical_issues.append(f"{issue_prefix} PDF CreationDate ({pdf_create_dt}) is significantly after Filesystem Modified ({fs_modification_dt.strftime('%Y-%m-%d %H:%M:%S') if fs_modification_dt else 'N/A'})")

                page_count = len(reader.pages)
                QTreeWidgetItem(pdf_root, ["Page Count", str(page_count)])


                if reader.is_encrypted:
                    sec_item = QTreeWidgetItem(pdf_root, ["🔒 ENCRYPTION", "Document is encrypted"])
                    sec_item.setForeground(0, QColor("red"))
                    self.anomalies.append(f"File '{base_name}': Encrypted PDF document")

                    if reader.metadata is None and page_count > 0:
                        meta_item = QTreeWidgetItem(pdf_root, ["⚠️ HIDDEN METADATA", "Metadata likely encrypted"])
                        meta_item.setForeground(0, QColor("orange"))
                        self.anomalies.append(f"File '{base_name}': PDF metadata likely hidden by encryption")


                producer = pdf_values.get("Producer", "").lower()
                if producer:
                    suspicious_producers = ["crack", "keygen", "patch", "converter"]
                    if any(bad in producer for bad in suspicious_producers):
                        prod_item = QTreeWidgetItem(pdf_root, ["🚨 SUSPICIOUS PRODUCER", pdf_values["Producer"]])
                        prod_item.setForeground(0, QColor("red"))
                        self.anomalies.append(f"File '{base_name}': Suspicious PDF producer - {pdf_values['Producer']}")

                if not has_pdf_content and page_count == 0:
                    QTreeWidgetItem(pdf_root, ["Status", "No metadata found and 0 pages."])
                    self.logical_issues.append(f"File '{base_name}': PDF has no metadata and 0 pages (potentially empty or corrupt).")
                elif not has_pdf_content:
                    QTreeWidgetItem(pdf_root, ["Status", "No standard metadata found."])

        except FileNotFoundError:
            self.anomalies.append(f"{base_name}: File not found during PDF processing.")
            QTreeWidgetItem(pdf_root, ["Error", "File not found."]).setForeground(1, QColor("red"))
        except PyPDF2.errors.PdfReadError as pdf_err:
            self.anomalies.append(f"{base_name}: Error reading PDF (likely corrupt or password protected) - {pdf_err}")
            QTreeWidgetItem(pdf_root, ["Error", f"Failed to read PDF: {pdf_err}"]).setForeground(1, QColor("red"))
        except Exception as e:
            err_type = type(e).__name__
            self.anomalies.append(f"{base_name}: Error processing PDF ({err_type}) - {e}")
            QTreeWidgetItem(pdf_root, ["Error", f"PDF processing failed: {e}"]).setForeground(1, QColor("red"))

    def process_docx(self, file_path, parent_item, stats):
        base_name = os.path.basename(file_path)
        if not Document:
            QTreeWidgetItem(parent_item, ["DOCX Status", "python-docx library missing."]).setForeground(1, QColor("orange"))
            return

        docx_root = QTreeWidgetItem(parent_item, ["DOCX Metadata"])
        has_docx_content = False

        try:
            doc = Document(file_path)
            core_properties = doc.core_properties


            properties = {
                "Title": core_properties.title,
                "Subject": core_properties.subject,
                "Author": core_properties.author,
                "Last Modified By": core_properties.last_modified_by,
                "Created": core_properties.created,
                "Modified": core_properties.modified,
                "Revision": core_properties.revision,
                "Category": core_properties.category,
                "Keywords": core_properties.keywords,
                "Comments": core_properties.comments,
            }

            docx_create_dt, docx_mod_dt = None, None
            for prop_name, prop_value in properties.items():
                if prop_value:
                    has_docx_content = True
                    value_str = str(prop_value)
                    if isinstance(prop_value, datetime):
                        try:
                            prop_value_aware = prop_value.astimezone()
                            value_str = prop_value_aware.strftime("%Y-%m-%d %H:%M:%S %Z%z")
                        except ValueError:
                            value_str = prop_value.strftime("%Y-%m-%d %H:%M:%S (Naive)")

                        if prop_name == "Created":
                            docx_create_dt = prop_value
                        if prop_name == "Modified":
                            docx_mod_dt = prop_value


                    if prop_name == "Last Modified By":
                        modifier_item = QTreeWidgetItem(docx_root, [prop_name, value_str])
                        last_mod = value_str.lower()
                        if "admin" in last_mod:
                            modifier_item.setForeground(1, QColor("orange"))
                            self.anomalies.append(f"File '{base_name}': Modified by admin account - {value_str}")
                        elif "temp" in last_mod or "user" in last_mod:
                            modifier_item.setForeground(1, QColor("blue"))
                            self.logical_issues.append(f"File '{base_name}': Modified by generic account - {value_str}")
                    else:
                        QTreeWidgetItem(docx_root, [prop_name, value_str])


            fs_creation_dt, fs_modification_dt = None, None
            if stats:
                try:
                    fs_modification_dt = datetime.fromtimestamp(stats.st_mtime)
                    ts_cre = stats.st_ctime
                    if hasattr(stats, 'st_birthtime') and stats.st_birthtime:
                        ts_cre = stats.st_birthtime
                    fs_creation_dt = datetime.fromtimestamp(ts_cre)
                except Exception as fs_e:
                    print(f"Warning: Could not extract filesystem times for comparison in {base_name}: {fs_e}")


            fs_mod_naive = fs_modification_dt.replace(tzinfo=None) if fs_modification_dt else None
            fs_cre_naive = fs_creation_dt.replace(tzinfo=None) if fs_creation_dt else None
            docx_cre_naive = docx_create_dt.replace(tzinfo=None) if docx_create_dt and docx_create_dt.tzinfo else docx_create_dt
            docx_mod_naive = docx_mod_dt.replace(tzinfo=None) if docx_mod_dt and docx_mod_dt.tzinfo else docx_mod_dt

            if docx_cre_naive or docx_mod_naive or fs_cre_naive or fs_mod_naive:
                issue_prefix = f"File '{base_name}' - DOCX Time Issue:"
                if docx_cre_naive and docx_mod_naive and docx_mod_naive < docx_cre_naive:
                    time_item = QTreeWidgetItem(docx_root, ["⚠️ TIME INCONSISTENCY", "Modified before Created"])
                    time_item.setForeground(0, QColor("orange"))
                    self.logical_issues.append(f"{issue_prefix} DOCX Modified ({docx_mod_dt}) is before DOCX Created ({docx_create_dt})")
                if docx_cre_naive and fs_mod_naive and docx_cre_naive > fs_mod_naive + timedelta(minutes=1):
                    time_item = QTreeWidgetItem(docx_root, ["⚠️ TIME INCONSISTENCY", "Created after filesystem modified"])
                    time_item.setForeground(0, QColor("orange"))
                    self.logical_issues.append(f"{issue_prefix} DOCX Created ({docx_create_dt}) is significantly after Filesystem Modified ({fs_modification_dt.strftime('%Y-%m-%d %H:%M:%S') if fs_modification_dt else 'N/A'})")


            stats_root = QTreeWidgetItem(docx_root, ["Document Statistics"])
            QTreeWidgetItem(stats_root, ["Paragraphs", str(len(doc.paragraphs))])
            QTreeWidgetItem(stats_root, ["Tables", str(len(doc.tables))])


            inline_shape_count = len(doc.inline_shapes)
            QTreeWidgetItem(stats_root, ["Inline Shapes (Images, etc.)", str(inline_shape_count)])
            if inline_shape_count > 20:
                shape_item = QTreeWidgetItem(stats_root, ["⚠️ MANY EMBEDDED OBJECTS", str(inline_shape_count)])
                shape_item.setForeground(0, QColor("orange"))
                self.logical_issues.append(f"File '{base_name}': Contains many embedded objects ({inline_shape_count})")


            if hasattr(doc, 'part') and hasattr(doc.part, 'rels'):
                for rel in doc.part.rels.values():
                    if 'vbaProject' in str(rel.target_ref):
                        macro_item = QTreeWidgetItem(docx_root, ["🚨 MACRO DETECTED", "Document contains VBA macros"])
                        macro_item.setForeground(0, QColor("red"))
                        self.anomalies.append(f"File '{base_name}': Contains VBA macros (potential security risk)")
                        break

            if not has_docx_content:
                QTreeWidgetItem(docx_root, ["Status", "No standard metadata found."])

        except FileNotFoundError:
            self.anomalies.append(f"{base_name}: File not found during DOCX processing.")
            QTreeWidgetItem(docx_root, ["Error", "File not found."]).setForeground(1, QColor("red"))
        except Exception as e:
            err_type = type(e).__name__
            if "zipfile.BadZipFile" in str(type(e)):
                self.anomalies.append(f"{base_name}: Error processing DOCX - File may be corrupt or not a valid DOCX (BadZipFile).")
                QTreeWidgetItem(docx_root, ["Error", "DOCX processing failed: BadZipFile (corrupt?)"]).setForeground(1, QColor("red"))
            else:
                self.anomalies.append(f"{base_name}: Error processing DOCX ({err_type}) - {e}")
                QTreeWidgetItem(docx_root, ["Error", f"DOCX processing failed: {e}"]).setForeground(1, QColor("red"))

    def process_media(self, file_path, parent_item, stats):
        base_name = os.path.basename(file_path)
        if not mutagen:
            QTreeWidgetItem(parent_item, ["Media Status", "mutagen library missing."]).setForeground(1, QColor("orange"))
            return

        media_root = QTreeWidgetItem(parent_item, ["Media Metadata"])
        has_media_content = False

        try:

            audio = mutagen.File(file_path, easy=True)
            if audio is None:

                 audio_complex = mutagen.File(file_path)
                 if audio_complex is None:
                     QTreeWidgetItem(media_root, ["Status", "File is not a supported media format or is corrupt."])
                     return
                 else:


                      QTreeWidgetItem(media_root, ["Status", "Standard tags not found, complex tags may exist."])

                      return


            common_tags = {
                'title': 'Title',
                'artist': 'Artist',
                'album': 'Album',
                'date': 'Date',
                'genre': 'Genre',
                'tracknumber': 'Track Number',
                'comment': 'Comment',
                'albumartist': 'Album Artist',
                'composer': 'Composer',
                'discnumber': 'Disc Number',
                'organization': 'Organization',
                'encodedby': 'Encoded By',
                'copyright': 'Copyright',
            }

            for tag_key, display_name in common_tags.items():
                if tag_key in audio:
                    has_media_content = True
                    value = audio[tag_key]

                    if isinstance(value, list):
                        value_str = ", ".join(str(v) for v in value)
                    else:
                         value_str = str(value)
                    QTreeWidgetItem(media_root, [display_name, value_str.strip()])


            if hasattr(audio, 'info'):
                tech_info_root = QTreeWidgetItem(media_root, ["Technical Info"])
                has_media_content = True
                info = audio.info
                tech_attrs = {
                    'length': 'Duration (s)',
                    'bitrate': 'Bitrate (bps)',
                    'sample_rate': 'Sample Rate (Hz)',
                    'channels': 'Channels',

                }
                for attr, display_name in tech_attrs.items():
                    if hasattr(info, attr):
                         value = getattr(info, attr)
                         if attr == 'length':
                             value = f"{value:.2f}"
                         elif attr == 'bitrate':
                              value = f"{value}"
                         QTreeWidgetItem(tech_info_root, [display_name, str(value)])


                if hasattr(audio, 'mime') and audio.mime:
                    QTreeWidgetItem(tech_info_root, ["MIME Type", ", ".join(audio.mime)])

                if hasattr(info, 'codec'):
                    QTreeWidgetItem(tech_info_root, ["Codec", str(info.codec)])


            if not has_media_content:
                QTreeWidgetItem(media_root, ["Status", "No standard metadata tags found."])

        except FileNotFoundError:
            self.anomalies.append(f"{base_name}: File not found during Media processing.")
            QTreeWidgetItem(media_root, ["Error", "File not found."]).setForeground(1, QColor("red"))
        except mutagen.MutagenError as e:
             self.anomalies.append(f"{base_name}: Error processing Media file (mutagen error) - {e}")
             QTreeWidgetItem(media_root, ["Error", f"Mutagen processing failed: {e}"]).setForeground(1, QColor("red"))
        except Exception as e:
            err_type = type(e).__name__
            self.anomalies.append(f"{base_name}: Error processing Media file ({err_type}) - {e}")
            QTreeWidgetItem(media_root, ["Error", f"Media processing failed: {e}"]).setForeground(1, QColor("red"))

    def show_anomalies(self):
        if self.anomalies:
            dialog = AnomaliesDialog(self.anomalies, self)
            dialog.exec()
        else:
            self.status_bar.showMessage("No anomalies detected.")

    def show_logical_issues(self):
        if self.logical_issues:
            dialog = LogicalIssuesDialog(self.logical_issues, self)
            dialog.exec()
        else:
            self.status_bar.showMessage("No logical issues detected.")

    def clear_results_ui_only(self):
        self.tree.clear()
        self.preview_widget.clear_preview()
        self.anomalies = []
        self.logical_issues = []
        self.anomalies_btn.setEnabled(False)
        self.anomalies_btn.setText(" Show Anomalies")
        self.logical_issues_btn.setEnabled(False)
        self.logical_issues_btn.setText(" Show Logical Issues")
        self.search_box.clear()

    def clear_results(self):
        self.file_paths = []
        if hasattr(self, 'tree'):
             self.tree.file_path_map = {}
        self.clear_results_ui_only()
        self.status_bar.showMessage("Ready")


if __name__ == '__main__':
    app = QApplication(sys.argv)


    styles = QStyleFactory.keys()
    if 'Fusion' in styles:
        app.setStyle('Fusion')
    elif 'WindowsVista' in styles:
        app.setStyle('WindowsVista')


    if platform.system() == 'Windows':

        try:
            import ctypes
            myappid = 'mycompany.myproduct.subproduct.version'
            ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)
        except Exception as e:
            print(f"Could not set AppUserModelID: {e}")

    main_window = MetadataAnalyzerApp()
    main_window.show()
    sys.exit(app.exec())